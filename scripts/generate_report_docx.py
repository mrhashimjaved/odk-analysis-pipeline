import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_report import OUTPUT_DIR, build_report_lines


# ── Output file paths ─────────────────────────────────────────────────────────
REPORT_6_MONTH_DOCX = os.path.join(OUTPUT_DIR, "6_month_follow_up_report.docx")
REPORT_9_MONTH_DOCX = os.path.join(OUTPUT_DIR, "9_month_endpoint_report.docx")
REPORT_CONSOLIDATED_DOCX = os.path.join(OUTPUT_DIR, "consolidated_report.docx")

REPORT_CONFIGS = [
    {
        "time_point": "6_month",
        "output_path": REPORT_6_MONTH_DOCX,
        "header_text": "SMART STEP — 6-Month Follow-Up Report",
    },
    {
        "time_point": "9_month",
        "output_path": REPORT_9_MONTH_DOCX,
        "header_text": "SMART STEP — 9-Month Endpoint Report",
    },
]


# ── Document configuration ────────────────────────────────────────────────────

def configure_document(document, header_text="ODK Analysis Pipeline Report"):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    title_style = document.styles["Title"]
    title_style.font.name = "Arial"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = header_text
    header.style = document.styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header.runs:
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ── Parsing helpers ───────────────────────────────────────────────────────────

def is_rule(line, char):
    return line and set(line) == {char}


def parse_table(lines, start_index):
    header = [cell.strip() for cell in lines[start_index].split("|")]
    rows = []
    index = start_index + 2
    while index < len(lines) and "|" in lines[index]:
        rows.append([cell.strip() for cell in lines[index].split("|")])
        index += 1
    return header, rows, index


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)
    run.font.bold = bold


def add_table(document, header, rows):
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for index, text in enumerate(header):
        set_cell_text(table.rows[0].cells[index], text, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            if index < len(cells):
                set_cell_text(cells[index], text)

    document.add_paragraph("")


def add_paragraph(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    return paragraph


# ── Builder ───────────────────────────────────────────────────────────────────

def build_docx(lines, output_path, header_text="ODK Analysis Pipeline Report"):
    document = Document()
    configure_document(document, header_text=header_text)

    index = 0
    first_title = True
    while index < len(lines):
        line = lines[index]
        next_line = lines[index + 1] if index + 1 < len(lines) else ""

        if not line.strip():
            index += 1
            continue

        if "|" in line and index + 1 < len(lines) and set(lines[index + 1].replace("|", "").strip()) == {"-"}:
            header, rows, index = parse_table(lines, index)
            add_table(document, header, rows)
            continue

        if is_rule(next_line, "="):
            paragraph = document.add_paragraph(line)
            paragraph.style = document.styles["Title" if first_title else "Heading 1"]
            paragraph.paragraph_format.space_after = Pt(10)
            first_title = False
            index += 2
            continue

        if is_rule(next_line, "-"):
            paragraph = document.add_paragraph(line)
            paragraph.style = document.styles["Heading 2"]
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            index += 2
            continue

        add_paragraph(document, line)
        index += 1

    document.save(output_path)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    any_written = False
    for config in REPORT_CONFIGS:
        time_point = config["time_point"]
        output_path = config["output_path"]
        header_text = config["header_text"]

        lines = build_report_lines(time_point=time_point)

        # Skip if there is no actual data (all sections empty except the title)
        data_lines = [l for l in lines if l.strip() and not set(l.strip()) <= {"=", "-"}]
        # Allow heading + intro, but skip if fewer than ~5 meaningful lines
        if len(data_lines) < 5:
            print(f"Skipping {output_path} — no data available for {time_point}.")
            continue

        build_docx(lines, output_path, header_text=header_text)
        print(f"Wrote {output_path}")
        any_written = True

    # Also write consolidated (all time points) for backward compatibility
    all_lines = build_report_lines(time_point=None)
    build_docx(all_lines, REPORT_CONSOLIDATED_DOCX, header_text="SMART STEP — Consolidated Report")
    print(f"Wrote {REPORT_CONSOLIDATED_DOCX}")


if __name__ == "__main__":
    main()
